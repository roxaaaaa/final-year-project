"""Build PDF and Word documents from a generated exam (topic, level, question strings)."""
from __future__ import annotations

import io
import os
import re
from typing import List

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _level_label(level: str) -> str:
    """Title case label for the paper level."""
    return "Higher" if (level or "").lower() == "higher" else "Ordinary"


def _escape_xml_text(s: str) -> str:
    """Escape &, <, > for ReportLab Paragraph XML."""
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_exam_docx(topic: str, level: str, questions: List[str]) -> bytes:
    """Build a simple .docx exam sheet; return raw bytes."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font_name = "Times New Roman"
    style.font_size = Pt(12)

    doc.add_heading("Leaving Certificate Practice", 0)
    doc.add_heading("Agricultural Science — " + _level_label(level), level=1)
    doc.add_paragraph(f"Topic: {topic or '—'}")
    doc.add_paragraph()

    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q or ''}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _register_unicode_font() -> str:
    """
    Register a TTF with good Latin-1 / Western European coverage (Irish fadas).
    ReportLab wheels may ship Vera.ttf but not DejaVu; avoid re-registering on repeat exports.
    """
    import reportlab

    fonts_dir = os.path.join(os.path.dirname(reportlab.__file__), "fonts")
    candidates = (
        ("DejaVuSans", "DejaVuSans.ttf"),
        ("Vera", "Vera.ttf"),
    )
    for register_as, filename in candidates:
        if register_as in pdfmetrics.getRegisteredFontNames():
            return register_as
        path = os.path.join(fonts_dir, filename)
        if not os.path.isfile(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(register_as, path))
            return register_as
        except Exception:
            continue
    return "Helvetica"


def build_exam_pdf(topic: str, level: str, questions: List[str]) -> bytes:
    """Build an A4 PDF with numbered questions; return raw bytes."""
    font_name = _register_unicode_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExamTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        leading=20,
        spaceAfter=6,
        alignment=1,
    )
    subtitle_style = ParagraphStyle(
        "ExamSubtitle",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceAfter=12,
        alignment=1,
    )
    body_style = ParagraphStyle(
        "ExamBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=14,
        spaceAfter=10,
    )
    topic_style = ParagraphStyle(
        "ExamTopic",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=14,
        spaceAfter=18,
        alignment=1,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    story = []
    story.append(Paragraph(_escape_xml_text("Leaving Certificate Practice"), title_style))
    story.append(
        Paragraph(
            _escape_xml_text(f"Agricultural Science — {_level_label(level)}"),
            subtitle_style,
        )
    )
    story.append(Paragraph(_escape_xml_text(f"Topic: {topic or '—'}"), topic_style))
    story.append(Spacer(1, 0.2 * cm))

    for i, q in enumerate(questions, start=1):
        text = _escape_xml_text(f"{i}. {q or ''}")
        story.append(Paragraph(text.replace("\n", "<br/>"), body_style))

    doc.build(story)
    return buf.getvalue()


def sanitize_download_filename(topic: str, exam_id: int, ext: str) -> str:
    """ASCII-safe filename for Content-Disposition."""
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", (topic or "exam").strip()).strip("-").lower()
    if not slug:
        slug = "exam"
    slug = slug[:60]
    safe_ext = "pdf" if ext.lower() == "pdf" else "docx"
    return f"agriexam-{exam_id}-{slug}.{safe_ext}"
